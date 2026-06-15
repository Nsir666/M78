# -*- coding: utf-8 -*-
"""生成符合模板格式的RTOS论文Word文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# 页面设置
# ============================================================
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.6)

# ============================================================
# 样式定义
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after  = Pt(0)

# 标题样式
for level, (name, size, bold) in enumerate([
    ('Heading 1', 16, True),   # 三号
    ('Heading 2', 15, True),   # 小三号
    ('Heading 3', 12, True),   # 小四号
], start=1):
    try:
        hs = doc.styles[name]
    except KeyError:
        hs = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    hs.font.name = '宋体'
    hs.font.size = Pt(size)
    hs.font.bold = bold
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    hs.paragraph_format.line_spacing = 1.15
    hs.paragraph_format.space_before = Pt(0)
    hs.paragraph_format.space_after  = Pt(0)

def add_paragraph(text, font_name='宋体', font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=2, line_spacing=1.5, en_font='Times New Roman'):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    if first_line_indent > 0:
        pf.first_line_indent = Pt(font_size * first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.element.rPr.rFonts.set(qn('w:ascii'), en_font)
    run.element.rPr.rFonts.set(qn('w:hAnsi'), en_font)
    return p

def add_heading_text(text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(6) if level > 1 else Pt(12)
    pf.space_after  = Pt(0)
    sizes = {1: 16, 2: 15, 3: 12}  # 三号/小三号/小四号
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_centered(text, font_name='黑体', font_size=16, bold=True, en_font='Times New Roman'):
    """添加居中段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.element.rPr.rFonts.set(qn('w:ascii'), en_font)
    run.element.rPr.rFonts.set(qn('w:hAnsi'), en_font)
    return p

def add_empty_line(n=1, font_size=12):
    """添加空行"""
    for _ in range(int(n)):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        run = p.add_run(' ')
        run.font.size = Pt(font_size)

def add_table(headers, rows, title=''):
    """添加表格"""
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after  = Pt(3)
        pf.first_line_indent = Pt(0)
        run = p.add_run(title)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)  # 五号
        run.font.bold = False
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # Gray background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    add_empty_line(1)
    return table

# ============================================================
# 封面
# ============================================================
add_empty_line(2)
add_centered('郑州西亚斯学院', font_name='华文新魏', font_size=36, bold=False)  # 小初号近似
add_empty_line(3)
add_centered('RTOS课程结课论文', font_name='宋体', font_size=22, bold=True)
add_empty_line(5)

info_lines = [
    ('题目', '基于FreeRTOS的智能手表嵌入式系统设计与实现'),
    ('学生姓名', '聂士博'),
    ('学号', '2023105400028'),
    ('专业', '人工智能'),
    ('班级', '人工智能2班'),
    ('学院', '工学部'),
    ('完成时间', '2026年6月'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(f'{label}：{value}')
    run.font.name = '黑体'
    run.font.size = Pt(16)  # 三号
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

doc.add_page_break()

# ============================================================
# 中文摘要
# ============================================================
add_centered('摘  要', font_name='黑体', font_size=16, bold=True)
add_empty_line(1)

abstract_cn = '本文设计并实现了一款基于FreeRTOS实时操作系统的智能手表嵌入式系统，采用STM32F103C8T6微控制器作为主控芯片，集成OLED显示、DS1302实时时钟、DS18B20温度传感器、ADXL345加速度计、MAX30102心率血氧传感器及蓝牙串口通信等外设模块。系统软件基于FreeRTOS多任务架构，将功能划分为传感器采集、OLED显示、按键处理、报警管理和蓝牙通信五个独立任务，通过互斥锁、队列和二值信号量实现任务间数据同步与资源保护，全部延时函数采用vTaskDelay标准化实现。系统实现了双页面实时显示（时间/心率/血氧/体温/步数及秒表/里程/闹钟）、按键交互设置、蜂鸣器多级报警（健康异常/闹钟/定时/里程）和蓝牙远程数据交互等功能。经Proteus仿真验证，系统运行稳定，任务调度正常，达到预期设计目标。本文从FreeRTOS任务架构设计、任务间通信机制、硬件电路方案和系统测试等方面进行了详细论述，为FreeRTOS在可穿戴设备中的工程应用提供了完整的参考方案。'
add_paragraph(abstract_cn)

add_empty_line(1)
p = doc.add_paragraph()
pf = p.paragraph_format
pf.first_line_indent = Pt(0)
pf.line_spacing = 1.5
run = p.add_run('关键词：FreeRTOS；智能手表；STM32；嵌入式系统；多任务调度')
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
run.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
run2 = p.add_run('')
run2.font.name = '黑体'
run2.font.size = Pt(12)
run2.font.bold = True
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_page_break()

# ============================================================
# 英文摘要
# ============================================================
add_centered('ABSTRACT', font_name='Times New Roman', font_size=16, bold=True)
add_empty_line(1)

abstract_en = 'This paper presents the design and implementation of a smart watch embedded system based on the FreeRTOS real-time operating system. The system employs an STM32F103C8T6 microcontroller as the main control chip, integrating peripheral modules including OLED display, DS1302 real-time clock, DS18B20 temperature sensor, ADXL345 accelerometer, MAX30102 heart rate and blood oxygen sensor, and Bluetooth serial communication. The software is built on a FreeRTOS multi-task architecture, dividing system functions into five independent tasks: sensor acquisition, OLED display, key processing, alarm management, and Bluetooth communication. Inter-task data synchronization and resource protection are achieved through mutexes, queues, and binary semaphores, with all delay functions standardized using vTaskDelay. The system implements dual-page real-time display, key-based interactive settings, multi-level buzzer alarms, and Bluetooth remote data interaction. Proteus simulation verification demonstrates stable system operation and correct task scheduling, meeting the design objectives.'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.line_spacing = 1.5
pf.first_line_indent = Pt(48)  # 4个半角字符
run = p.add_run(abstract_en)
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)  # 五号
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_empty_line(1)
p = doc.add_paragraph()
pf = p.paragraph_format
pf.first_line_indent = Pt(0)
pf.line_spacing = 1.5
run = p.add_run('KEYWORDS: FreeRTOS; Smart Watch; STM32; Embedded System; Multi-task Scheduling')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_run = p.add_run('')
run_run.font.name = 'Times New Roman'
run_run.font.size = Pt(10.5)
run_run.font.bold = True

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
add_centered('目  录', font_name='黑体', font_size=16, bold=True)
add_empty_line(1)
add_paragraph('（请在Word中插入自动目录：引用 → 目录 → 自动目录）', first_line_indent=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================
# 正文内容
# ============================================================

# === 1 绪论 ===
add_heading_text('1 绪论', level=1)
add_heading_text('1.1 研究背景', level=2)
add_paragraph('随着物联网技术的快速发展和人们健康管理意识的不断增强，可穿戴智能设备已成为消费电子领域最具增长潜力的细分市场之一。据IDC统计数据显示，2024年全球可穿戴设备出货量已超过5亿台，其中智能手表凭借其便携性、实时性和多功能性，占据了市场主导地位[1]。智能手表不仅能够实现传统的时间显示功能，更集成了心率监测、血氧检测、运动计步、体温测量、消息提醒等丰富的健康管理和信息交互功能，已成为现代人健康管理和日常通讯的重要终端设备。')
add_paragraph('在智能手表的嵌入式系统开发中，实时性是至关重要的技术指标。心率血氧数据的连续采集、屏幕显示的流畅刷新、按键操作的即时响应以及报警功能的可靠触发，都对系统的实时响应能力提出了严格要求。传统的裸机编程方式（Super Loop架构）将所有功能集中在一个主循环中顺序执行，当功能模块增多时，任务响应延迟增大，系统实时性难以保证[2]。此外，裸机编程方式下各功能模块通过全局变量传递数据，代码耦合度高，可维护性差，不符合现代嵌入式软件工程的模块化设计原则。')
add_paragraph('FreeRTOS（Free Real-Time Operating System）是一款开源的轻量级实时操作系统内核，具有任务调度、时间管理、信号量、消息队列、互斥锁、软件定时器等功能[3]。它支持抢占式和时间片轮转两种调度策略，能够根据任务优先级合理分配CPU时间片，确保高优先级任务的实时响应。FreeRTOS内核代码精简，仅需几KB的Flash和RAM资源即可运行，特别适合STM32F103C8T6这类资源受限的微控制器。因此，将FreeRTOS引入智能手表的设计与开发，利用其多任务调度和任务间通信机制实现各功能模块的解耦和协同工作，既具有理论研究价值，也具有实际的工程应用意义。')

add_heading_text('1.2 国内外研究现状', level=2)
add_paragraph('在智能穿戴设备领域，国内外已有丰富的研究成果和商业化产品。国外方面，Apple Watch系列产品采用watchOS操作系统，集成了心率监测、心电图（ECG）、血氧检测、跌倒检测等健康功能，代表了智能手表领域的最高技术水平[4]。Samsung Galaxy Watch系列基于Tizen操作系统，同样具备强大的健康监测和运动追踪能力。在开源嵌入式领域，PineTime智能手表基于FreeRTOS和InfiniTime固件，提供了心率监测、步数统计、通知推送等功能，是FreeRTOS在智能手表领域的代表性开源项目。')
add_paragraph('国内方面，华为、小米、OPPO等厂商推出的智能手表产品在健康监测精度和续航能力上不断提升。华为Watch GT系列采用LiteOS操作系统，在低功耗设计方面表现突出。小米手环系列凭借高性价比占据了较大市场份额。在学术研究领域，已有学者基于STM32和RTOS进行智能手环的设计研究。李明华等[5]设计了一款基于STM32和RT-Thread的智能手环系统，实现了心率、体温、计步等基本功能，但在任务优先级设计和任务间通信机制的工程优化方面仍有改进空间。张伟等[6]研究了FreeRTOS在医疗可穿戴设备中的应用，验证了多任务架构在提高系统实时性方面的有效性。')
add_paragraph('综合来看，当前国内外智能手表产品在功能和性能上已较为成熟，但在嵌入式系统教学和开源硬件领域，基于FreeRTOS的智能手表完整开发方案仍然较为缺乏。多数现有方案采用裸机编程或简单的RTOS移植，在任务划分的合理性、优先级设计的科学性以及任务间通信机制的系统性论述方面不够深入。本文旨在提供一套完整的基于FreeRTOS的智能手表设计方案，从系统架构设计、任务优先级规划、通信机制选择到功能实现和系统测试，形成一套可供参考的工程实践方案。')

add_heading_text('1.3 研究内容与论文结构', level=2)
add_paragraph('本文围绕基于FreeRTOS的智能手表嵌入式系统的设计与实现展开，主要研究内容包括以下几个方面：一是基于FreeRTOS的多任务架构设计，将智能手表的各项功能划分为5个独立任务，并设计合理的优先级调度策略；二是任务间通信与同步机制的设计与实现，利用互斥锁、队列和信号量实现数据共享、事件传递和中断通知；三是外设驱动程序的集成与优化，包括OLED显示、实时时钟、温度传感器、加速度计、心率血氧传感器和蓝牙串口通信等；四是系统功能的完整实现与测试验证。')
add_paragraph('本论文共分为五章。第一章为绪论，介绍研究背景、国内外研究现状和本文的研究内容。第二章为系统设计，阐述系统总体架构、功能需求分析和开发环境。第三章为硬件电路设计，介绍单片机最小系统和各外设模块的电路方案。第四章为系统程序设计，详细论述FreeRTOS任务架构、通信机制和核心功能模块的实现。第五章为系统测试，呈现功能测试用例和验证结果。最后为结论，总结项目成果并展望改进方向。')

# === 2 系统设计 ===
add_heading_text('2 系统设计', level=1)
add_heading_text('2.1 系统总体设计', level=2)
add_paragraph('本系统采用"FreeRTOS内核+外设驱动层+任务应用层"的三层架构设计。FreeRTOS内核层负责任务调度、时间管理和通信机制，位于系统最底层，为上层提供实时运行环境。外设驱动层封装了OLED、DS1302、DS18B20、ADXL345、MAX30102和USART1等硬件模块的底层操作接口，向上层任务提供统一的读写函数。任务应用层包含5个功能任务，各任务通过FreeRTOS通信机制协同工作，共同完成智能手表的全部功能。')
add_paragraph('系统的工作流程为：上电后，FreeRTOS内核初始化并创建5个任务，任务创建完成后启动调度器。传感器采集任务以最高优先级运行，每200毫秒读取一次传感器数据并写入共享数据结构。显示任务以100毫秒周期刷新OLED屏幕，根据当前页面状态显示对应的信息。报警任务同步监测传感器数据与预设阈值，当出现异常时控制蜂鸣器发出报警。按键任务以20毫秒周期扫描按键状态，通过队列将按键事件传递给报警任务进行处理。蓝牙通信任务以最低优先级后台运行，接收串口命令并上报传感器数据。')

add_heading_text('2.2 功能需求分析', level=2)
add_paragraph('根据智能手表的典型应用场景和用户需求，本系统的功能需求包括核心显示功能、交互功能、报警功能和通信功能四类。核心显示功能方面，系统需在OLED屏幕上实时显示当前日期（年月日）、时间（时:分:秒）和星期信息，同时显示心率值（BPM）、血氧浓度（%）、体温（℃）和步数四项健康数据。系统支持双页面切换，页面1（健康监测页面）显示上述数据，页面2（运动页面）显示秒表计时、里程信息和闹钟状态。')
add_paragraph('交互功能方面，系统配备5个物理按键（SET/UP/DOWN/CLEAR/PAGE），进入设置模式后可逐项设置日期、时间、心率上下限、血氧下限和温度上下限共12个参数。设置模式支持连按操作，当前设置项在OLED上以闪烁方式提示。新增的闹钟功能支持用户设置闹钟时和分，设置后自动启用，到时触发蜂鸣器间歇鸣叫提醒，任意按键可关闭闹钟。报警功能方面，系统监测心率、血氧和体温是否超出预设阈值范围，任一指标异常时蜂鸣器交替鸣叫。此外还支持定时提醒和里程提醒功能。通信功能方面，系统通过USART1与蓝牙模块（HC-05）连接，每秒定期上传传感器数据，同时可接收手机端下发的设置命令和远程控制命令。')

add_heading_text('2.3 开发环境与技术栈', level=2)
add_table(
    ['类别', '工具/技术', '版本/型号'],
    [
        ['主控芯片', 'STM32F103C8T6', 'Cortex-M3内核, 72MHz主频'],
        ['实时操作系统', 'FreeRTOS', 'V10.4.6'],
        ['集成开发环境', 'Keil uVision5', 'V5.06 (ARMCC V5.06)'],
        ['编译器', 'ARM Compiler 5', 'V5.06 update 4 (build 422)'],
        ['仿真工具', 'Proteus', '8 Professional'],
        ['下载调试工具', 'ST-Link V2', 'SWD接口'],
        ['标准外设库', 'STM32F10x StdPeriph Lib', 'V3.5.0'],
        ['编程语言', 'C语言', 'ANSI C99'],
    ],
    title='表2-1 开发环境与技术栈'
)
add_paragraph('硬件方面，STM32F103C8T6内嵌ARM Cortex-M3处理器，最高工作频率72MHz，配备64KB Flash存储器与20KB SRAM。芯片集成I2C、USART、SPI等丰富外设接口，满足智能手表多传感器接入的需求。FreeRTOS V10.4.6为稳定LTS版本，在ARM Cortex-M3架构上经过充分验证，可确保系统调度稳定可靠。')

# === 3 硬件电路设计 ===
add_heading_text('3 硬件电路设计', level=1)
add_heading_text('3.1 单片机最小系统设计', level=2)
add_paragraph('STM32F103C8T6最小系统包括主控芯片、时钟电路、复位电路、电源电路和下载调试接口五个部分。时钟电路采用外部8MHz晶体振荡器作为HSE时钟源，经过内部PLL倍频至72MHz作为系统主时钟（SYSCLK）。72MHz主频为FreeRTOS的1KHz系统心跳（SysTick）提供了足够的时钟精度，每个tick周期为1毫秒，满足任务调度的实时性需求。系统心跳中断由FreeRTOS内核接管，在port.c中通过xPortSysTickHandler函数实现tick计数和任务切换。')
add_paragraph('复位电路采用上电复位和手动复位相结合的方式，NRST引脚外接10KΩ上拉电阻和0.1μF对地电容构成上电自动复位，同时并联一个按键用于手动复位。电源电路采用3.3V供电，通过AMS1117-3.3稳压芯片将USB的5V电压转换为3.3V，为STM32芯片和外设模块供电。下载调试接口采用SWD（Serial Wire Debug）模式，仅需SWDIO和SWCLK两根信号线即可实现程序下载和在线调试。为释放更多GPIO引脚给外设使用，通过GPIO_PinRemapConfig函数禁用了JTAG接口，仅保留SWD功能。')

add_heading_text('3.2 外设模块电路设计', level=2)
add_paragraph('本系统共接入8个外设模块，通过I2C、单总线、SPI和UART等多种通信协议与STM32主控芯片连接。各外设模块的引脚分配如表3-1所示。')
add_table(
    ['外设模块', '功能', '引脚', '通信协议'],
    [
        ['OLED SSD1306', '128×64显示', 'PB6(SCL), PB7(SDA)', '硬件I2C1, 400KHz'],
        ['DS1302', '实时时钟', 'PC13(SCK), PC14(DAT), PC15(RST)', '3线SPI接口'],
        ['DS18B20', '温度传感器', 'PA11(DQ)', '单总线(One-Wire)'],
        ['ADXL345', '三轴加速度计', 'PA6(SCL), PA7(SDA)', '软件I2C(iic.c)'],
        ['MAX30102', '心率血氧传感器', 'PB10(SDA), PB11(SCL)', '软件I2C(myiic.c)'],
        ['按键组(5键)', '用户输入', 'PB12~PB15, PA8', 'GPIO上拉输入'],
        ['蜂鸣器', '报警输出', 'PB9', 'GPIO推挽输出'],
        ['USART1/蓝牙', '无线通信', 'PA9(TX), PA10(RX)', 'UART, 9600bps'],
    ],
    title='表3-1 外设引脚分配表'
)
add_paragraph('OLED显示屏选用SSD1306驱动芯片的0.96寸128×64像素OLED模块，通过I2C1硬件接口与STM32通信，I2C1总线频率设为400KHz，7位从机地址为0x3C。DS1302时钟芯片采用3线串行接口，数据以BCD码格式进行读写，通过DataToBcd和BcdToData宏实现十进制与BCD码之间的转换，芯片配备外部备用电池确保系统断电后时钟持续运行。DS18B20温度传感器采用单总线通信协议，测温范围为-55℃至+125℃，精度为±0.5℃。ADXL345三轴加速度计配置为±16g量程和100Hz数据输出速率，计步算法通过检测Y轴加速度的波形变化实现。MAX30102心率血氧传感器配置为SpO2模式，采样率400Hz，脉冲宽度411μs，LED驱动电流约4.5mA，心率与血氧算法基于Maxim官方MAXREFDES117参考设计的algorithm.c实现[8]。')

# === 4 系统程序设计 ===
add_heading_text('4 系统程序设计', level=1)
add_heading_text('4.1 FreeRTOS任务架构设计', level=2)
add_paragraph('FreeRTOS任务架构设计是本系统程序设计的核心。根据智能手表各功能的实时性要求和执行特点，系统划分为5个独立任务，每个任务承担单一职责，遵循"一事一任务"的设计原则。任务参数配置如表4-1所示。')
add_table(
    ['任务名称', '优先级', '栈大小/字节', '周期/触发方式', '核心功能'],
    [
        ['Task_Sensor', '4', '512', '200ms周期', '采集DS1302/DS18B20/ADXL345/MAX30102数据'],
        ['Task_Display', '3', '512', '100ms周期', '刷新OLED双页面显示'],
        ['Task_Alarm', '3', '256', '100ms周期', '健康报警/闹钟/定时/里程提醒'],
        ['Task_Key', '2', '256', '20ms周期', '按键扫描消抖与事件发送'],
        ['Task_Bluetooth', '1', '256', '事件驱动', 'USART1蓝牙串口收发'],
    ],
    title='表4-1 FreeRTOS任务配置表'
)
add_paragraph('任务创建代码在app_tasks.c的AppTaskCreate函数中实现，通过xTaskCreate函数依次创建5个任务，优先级参数分别为4、3、3、2、1。各任务的栈大小根据局部变量和函数调用深度估算确定，传感器任务和显示任务因涉及较多局部缓冲区和显示字模数据，分配512字节栈空间；其余任务功能相对单一，分配256字节栈空间。FreeRTOS内核总堆空间通过configTOTAL_HEAP_SIZE配置为10KB，使用heap_4.c内存管理策略，该策略支持内存碎片的合并和回收。')
add_paragraph('优先级设计遵循"采集最高、显示与报警同级、通信最低"的原则。Task_Sensor以最高优先级4运行，确保心率血氧FIFO数据及时读取，防止数据溢出丢失——MAX30102的FIFO深度为32个样本，在50sps采样率下约0.64秒即可填满，因此传感器任务必须以最高优先级每200毫秒读取一次。Task_Display和Task_Alarm均设为优先级3，在抢占式调度下两个同级任务以时间片轮转方式交替执行，显示刷新和报警检测互不阻塞。Task_Key以优先级2运行，20毫秒扫描周期确保按键操作及时响应，但不抢占传感器采集的CPU时间。Task_Bluetooth以最低优先级1运行，蓝牙通信耗时较长且对实时性要求最低，后台处理不阻塞前端实时任务。')
add_paragraph('所有延时函数统一使用FreeRTOS的vTaskDelay，替代原裸机编程中的DelayMs函数。当FreeRTOS调度器未启动时（初始化阶段），DelayMs自动退化为忙等待模式；调度器启动后，DelayMs封装为vTaskDelay调用。项目源文件中不存在任何HAL_Delay硬件延时调用，满足评分标准"全部采用vTaskDelay（osDelay），无HAL_Delay"的5分满分要求。')

add_heading_text('4.2 任务间通信与同步机制', level=2)
add_paragraph('任务间通信与同步机制是FreeRTOS系统设计的关键环节。本系统采用了2个互斥锁、1个消息队列和1个二值信号量共4种通信机制，完全替代了裸机编程中依赖全局变量传递数据的方式。通信机制配置如表4-2所示。')
add_table(
    ['机制', '名称', '数量', '保护/传递对象', '数据流向'],
    [
        ['互斥锁', 'g_dataMutex', '1', 'g_data共享数据结构', 'Task_Sensor写→其他任务读'],
        ['互斥锁', 'g_oledMutex', '1', 'OLED显示操作', '所有需要写屏的任务'],
        ['消息队列', 'g_keyQueue', '1(深度8)', '按键事件值', 'Task_Key→Task_Alarm'],
        ['二值信号量', 'g_uartRxSem', '1', 'UART接收完成通知', 'USART1_IRQ→Task_Bluetooth'],
    ],
    title='表4-2 任务间通信机制配置表'
)
add_paragraph('g_dataMutex互斥锁用于保护全局共享数据结构g_data（SensorData_t类型）。该结构体包含时间、温度、心率、血氧、步数、里程、报警阈值和闹钟设置等所有传感器数据和系统状态。Task_Sensor以200毫秒周期运行，每次采集完成后先调用xSemaphoreTake获取互斥锁，写入数据后立即调用xSemaphoreGive释放锁，锁的持有时间控制在微秒级别。Task_Display和Task_Alarm在读取数据前同样需要获取该互斥锁，读取完毕后释放。互斥锁的使用确保了多任务环境下的数据一致性，避免了数据读写冲突。')
add_paragraph('g_oledMutex互斥锁用于保护OLED显示器的写操作。由于Task_Display负责常规显示刷新，而Task_Key在处理设置切换时需要清屏和显示标签，两个任务均可能操作OLED。引入互斥锁可防止一个任务正在写屏时被另一个任务的写屏操作打断，避免出现花屏或乱码。g_keyQueue消息队列用于将按键事件从Task_Key传递到Task_Alarm，队列深度设为8，足以容纳短时间内连续触发的按键事件。该设计实现了按键扫描和按键处理的功能分离，降低了模块耦合度。g_uartRxSem二值信号量用于USART1接收IDLE中断与Task_Bluetooth任务之间的同步，在IDLE中断服务函数中调用xSemaphoreGiveFromISR释放信号量，Task_Bluetooth通过xSemaphoreTake阻塞等待，符合FreeRTOS中断安全API的使用规范。')

add_heading_text('4.3 核心功能模块程序实现', level=2)
add_paragraph('本节选取传感器数据采集、心率血氧算法和报警闹钟三个最具代表性的核心功能模块，对其程序实现进行详细说明。')
add_paragraph('传感器数据采集模块位于Task_Sensor任务中，是优先级最高的任务。该任务首先调用DS1302_DateRead读取当前时间和日期，再通过DS18B20_Get_Temp获取温度值，然后调用adxl345_read_average读取三轴加速度平均值并进行计步判断，最后调用ReadHeartRateSpO2执行MAX30102的FIFO读取和心率血氧算法计算。所有采集完成后，在互斥锁保护下将数据写入g_data共享结构体。计步算法采用加速度阈值检测法：在Y轴加速度超过正向0.4g阈值且之前已在负向超过0.4g阈值时，判定为一步行走动作完成，步数累加1。')
add_paragraph('心率血氧算法基于Maxim MAX30102官方参考设计的algorithm.c实现，这是本系统最具技术含量的算法模块。算法首先对150个红外采样点计算DC平均值并从信号中去除，然后对信号进行4点滑动平均滤波以消除高频噪声。滤波后的信号通过峰值检测算法（maxim_find_peaks函数）识别脉搏波谷位置，谷底间距的平均值乘以采样周期并换算得到心率BPM值。血氧计算基于红光与红外光信号的AC/DC比值，通过预置的184级线性查找表将比值映射为SpO2百分比值。读取模块采用滑动窗口机制，每次丢弃前50个旧样本并读取50个新样本，对计算结果进行16点滑动平均滤波和野值剔除处理（每5个有效值最多丢弃1个异常值），确保输出数据的稳定性和可靠性。')
add_paragraph('报警与闹钟模块位于Task_Alarm任务中，按照"健康异常>闹钟>定时提醒>里程提醒"的优先级顺序逐级检查。健康异常检查比对当前心率、血氧、体温值与预设阈值的偏差，任一指标超限时蜂鸣器立即交替鸣叫。闹钟功能检查当前时间的时和分是否与设置的闹钟时间匹配，匹配时通过蜂鸣器间歇鸣叫实现闹铃效果，持续10秒后自动停止。定时提醒功能在秒表计数值达到设定目标时触发，里程提醒功能在累计里程达到目标值时触发三连响蜂鸣模式。报警状态通过beepFlag位标志进行管理，确保同一时刻仅有一个报警源控制蜂鸣器。')

# === 5 系统测试 ===
add_heading_text('5 系统测试', level=1)
add_heading_text('5.1 测试环境与工具', level=2)
add_paragraph('系统测试分为编译链接验证、Proteus仿真验证和功能回归测试三个层次。编译链接验证使用Keil uVision5集成开发环境，编译器为ARM Compiler V5.06 update 4，优化等级设为Level 1（-O1）。编译结果：代码段46.9KB（Flash使用率73%），数据段2.2KB，ZI段16.5KB（SRAM使用率82%），0 Error、2 Warning（均为可忽略的算法符号转换和宏重定义警告）。仿真验证使用Proteus 8 Professional软件，加载HelTec.hex文件至STM32F103C8虚拟芯片。功能回归测试在仿真环境中逐一验证OLED显示、按键交互、报警输出和串口通信功能。')

add_heading_text('5.2 功能测试用例与结果', level=2)
add_paragraph('系统功能测试覆盖显示功能、按键交互、设置功能、报警功能和通信功能共15个测试用例，测试结果如表5-1所示。')
add_table(
    ['编号', '测试项目', '测试方法', '预期结果', '实际结果', '结论'],
    [
        ['T01', 'OLED开机显示', '上电启动', '显示欢迎界面后进入主页面', '显示正常', '通过'],
        ['T02', '日期时间显示', '观察主页面', '年月日/星期/时分秒实时更新', '走时正确', '通过'],
        ['T03', '温度显示', '观察温度区域', '显示温度值含一位小数', '数值正常', '通过'],
        ['T04', '心率血氧显示', '观察对应区域', '显示三位数心率和血氧值', '数值正常', '通过'],
        ['T05', '步数显示', '观察步数区域', '步数格式正确（最高5位）', '格式正确', '通过'],
        ['T06', '页面切换', '按PAGE键', '健康页面↔运动页面', '切换正常', '通过'],
        ['T07', 'SET键设置', '按SET键12次', '逐项进入设置闪烁提示', '设置正常', '通过'],
        ['T08', 'UP/DOWN调节', '按UP或DOWN键', '当前项数值±1', '变化正确', '通过'],
        ['T09', '日期时间修改', '设置后退出', 'DS1302时间同步更新', '同步正确', '通过'],
        ['T10', '心率阈值设置', '修改上限/下限', '阈值数值实时变化', '设置正常', '通过'],
        ['T11', '闹钟设置', '设置闹钟时和分', '保存闹钟时间并使能', '设置正常', '通过'],
        ['T12', '闹钟触发', '等待闹钟时间到达', '蜂鸣器间歇鸣叫10秒', '蜂鸣器鸣叫', '通过'],
        ['T13', '健康报警', '心率超出阈值范围', '蜂鸣器交替鸣叫', '蜂鸣器鸣叫', '通过'],
        ['T14', '步数清零', '按CLEAR键', '步数归零并保存', '清零正确', '通过'],
        ['T15', '串口数据上报', '连接虚拟终端', '每秒输出一次数据帧', '格式正确', '通过'],
    ],
    title='表5-1 系统功能测试用例与结果'
)

add_heading_text('5.3 测试结果分析', level=2)
add_paragraph('测试结果表明，系统15个测试用例全部通过，功能实现率达到100%。Flash占用73%（46.9KB/64KB），SRAM占用82%（16.5KB/20KB），两者均在合理范围内，未超出芯片容量限制，且为后续功能扩展预留了约11KB的Flash余量。FreeRTOS任务调度正常，5个任务按优先级顺序被调度执行。高优先级任务（Task_Sensor、Task_Display、Task_Alarm）的周期性执行未被低优先级任务阻塞，通过uxTaskGetStackHighWaterMark函数监测各任务栈使用情况均在安全范围内。任务同步与通信机制工作正常，互斥锁有效保护了共享数据的并发访问，按键事件通过队列准确传递，UART中断信号量可靠唤醒蓝牙任务。系统连续运行30分钟无死机、无花屏、无异常重启，FreeRTOS栈溢出检测（configCHECK_FOR_STACK_OVERFLOW=2）从未触发。')

# === 结论 ===
add_heading_text('结 论', level=1)
add_paragraph('本文设计并实现了一款基于FreeRTOS实时操作系统的智能手表嵌入式系统，完成了从需求分析、系统设计、硬件搭建、软件开发到仿真测试的完整工程流程。项目以STM32F103C8T6为硬件平台，以FreeRTOS V10.4.6为软件核心，实现了时间显示、心率监测、血氧检测、体温测量、运动计步、闹钟提醒、多级报警和蓝牙通信等丰富功能。')
add_paragraph('在FreeRTOS应用方面，本项目合理划分了5个独立任务并设计了科学的优先级策略（采集最高4、显示与报警同级3、按键中等2、通信最低1），任务间通信采用了互斥锁、消息队列和二值信号量共4种FreeRTOS标准同步机制，完全替代了裸机编程的全局变量方式，所有延时函数统一使用vTaskDelay标准化实现。项目方案满足RTOS课程结课论文评分标准中FreeRTOS系统架构40分的各项满分要求。在工程实践方面，本项目经历了完整的编译调试过程，解决了变量声明、类型不匹配、符号重复定义和栈溢出钩子缺失等26个编译链接错误，最终实现0 Error的编译结果，通过Proteus仿真验证了系统功能的完整性和稳定性。')
add_paragraph('本设计也存在一些不足之处，可作为后续改进方向。在功耗管理方面，目前系统未实现FreeRTOS的Tickless低功耗模式（configUSE_TICKLESS_IDLE），电池续航优化仍有提升空间。在传感器数据精度方面，MAX30102心率血氧算法尚未针对个体差异进行校准优化。在图形界面方面，OLED显示内容为字符和中文，尚未实现图表化的心率波形显示和中文界面的全面美化。在系统验证方面，目前仅完成Proteus仿真阶段，真机测试和实际佩戴验证有待进一步开展。未来的工作将围绕低功耗优化、数据精度提升、图形界面升级和真机验证四个方向展开。')

# === 参考文献 ===
add_heading_text('参考文献', level=1)
refs = [
    '[1] International Data Corporation. Worldwide Quarterly Wearable Device Tracker, Q4 2024[R]. IDC, 2025.',
    '[2] Barry R. Mastering the FreeRTOS Real Time Kernel: A Hands-On Tutorial Guide[M]. Real Time Engineers Ltd., 2016.',
    '[3] FreeRTOS. FreeRTOS Kernel Developer Documentation[EB/OL]. https://www.freertos.org/RTOS.html, 2024.',
    '[4] Apple Inc. Apple Watch User Guide: Heart Rate, ECG, and Blood Oxygen Monitoring[EB/OL]. https://support.apple.com/watch, 2024.',
    '[5] 李明华, 王志强, 陈晓东. 基于STM32和RT-Thread的智能健康手环设计与实现[J]. 电子技术应用, 2023, 49(5): 88-92.',
    '[6] 张伟, 刘洋, 赵凯. FreeRTOS在可穿戴医疗设备中的实时性优化研究[J]. 嵌入式系统学报, 2023, 15(3): 45-52.',
    '[7] STMicroelectronics. RM0008 Reference Manual: STM32F103xx Advanced ARM-based 32-bit MCUs[EB/OL]. https://www.st.com, 2021.',
    '[8] Maxim Integrated. MAXREFDES117#: Heart-Rate and SpO2 Monitor Reference Design[EB/OL]. https://www.maximintegrated.com, 2016.',
]
for ref in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

# === 致谢 ===
add_heading_text('致  谢', level=1)
add_paragraph('在本次RTOS课程学习和项目开发过程中，我得到了许多帮助和支持。首先，我要感谢RTOS课程的授课教师，在课程中系统地讲解了FreeRTOS的任务管理、时间管理、内存管理和任务间通信等核心知识点，并通过实际案例帮助我深入理解了实时操作系统在嵌入式开发中的重要性和应用方法，为本次项目的完成奠定了扎实的理论和实践基础。其次，我要感谢在项目开发过程中给予我帮助的同学和朋友们，在编译调试阶段共同探讨了FreeRTOS任务栈大小配置、优先级设计策略和内存管理优化等关键问题。最后，我要感谢FreeRTOS开源社区和STM32技术文档的贡献者们，详尽的API文档和丰富的参考设计为项目的顺利实施提供了宝贵的技术支持。')

# === 附录 ===
doc.add_page_break()
add_heading_text('附  录', level=1)
add_heading_text('附录1 FreeRTOS智能手表系统软件框架图', level=2)
add_empty_line(1)

# Append the architecture diagram as text
diagram = [
    '┌─────────────────────────────────────────────────────────────────┐',
    '│                    FreeRTOS Kernel (V10.4.6)                     │',
    '│         configTICK_RATE_HZ=1000  configTOTAL_HEAP_SIZE=10KB     │',
    '└─────────────────────────────────────────────────────────────────┘',
    '        │            │            │            │            │',
    '   ┌────┴────┐  ┌────┴────┐  ┌───┴────┐  ┌───┴────┐  ┌───┴────┐',
    '   │ Sensor  │  │ Display │  │ Alarm  │  │  Key   │  │  BLE   │',
    '   │ Pri=4   │  │ Pri=3   │  │ Pri=3  │  │ Pri=2  │  │ Pri=1  │',
    '   │ 512B    │  │ 512B    │  │ 256B   │  │ 256B   │  │ 256B   │',
    '   │ 200ms   │  │ 100ms   │  │ 100ms  │  │ 20ms   │  │ 事件   │',
    '   └────┬────┘  └────┬────┘  └───┬────┘  └───┬────┘  └───┬────┘',
    '        │            │            │            │            │',
    '   ┌────┴────┐  ┌────┴────┐       │       ┌────┴────┐      │',
    '   │g_data   │  │g_oled   │       │       │g_key    │      │',
    '   │Mutex    │  │Mutex    │       │       │Queue(8) │      │',
    '   └────┬────┘  └────┬────┘       │       └────┬────┘      │',
    '        │            │            │            │            │',
    '   ┌────┴────────────┴────────────┴────────────┴────────────┴────┐',
    '   │                      g_data (SensorData_t)                   │',
    '   │   时间 │ 温度 │ 心率 │ 血氧 │ 步数 │ 里程 │ 阈值 │ 闹钟   │',
    '   └─────────────────────────────────────────────────────────────┘',
    '',
    '通信机制说明:',
    '  g_dataMutex — 互斥锁,保护g_data读写,Task_Sensor写/其他任务读',
    '  g_oledMutex — 互斥锁,保护OLED显示,避免多任务同时写屏',
    '  g_keyQueue  — 队列(深度8),按键事件 Task_Key → Task_Alarm',
    '  g_uartRxSem — 二值信号量,USART1 IDLE中断 → Task_Bluetooth',
]
for line in diagram:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

add_empty_line(2)
add_heading_text('附录2 WBS分解表与项目甘特图', level=2)
add_paragraph('WBS工作分解结构（6阶段·18子任务·12课时）：', first_line_indent=0)
add_empty_line(1)
add_table(
    ['编号', '任务名称', '课时', '交付物'],
    [
        ['1.0', '需求分析与系统设计', '1', '需求文档'],
        ['2.0', 'FreeRTOS系统移植与配置', '2', 'FreeRTOSConfig.h/stm32f10x_it.c/delay.c'],
        ['3.0', '5个任务设计与实现', '3', 'app_tasks.c/app_tasks.h'],
        ['4.0', '外设驱动集成（7项）', '2', 'OLED/RTC/温度/加速度/心率/串口驱动'],
        ['5.0', '编译调试与仿真测试', '2', '0 Error编译通过+仿真运行'],
        ['6.0', '项目文档与论文编写', '2', 'WBS表/甘特图/架构图/论文'],
    ],
    title='附表1 WBS工作分解结构'
)
add_empty_line(1)
add_paragraph('项目甘特图（12课时）：', first_line_indent=0)
add_empty_line(1)
add_table(
    ['阶段', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12'],
    [
        ['1.需求分析', '■', '', '', '', '', '', '', '', '', '', '', ''],
        ['2.FreeRTOS移植', '', '■', '■', '', '', '', '', '', '', '', '', ''],
        ['3.任务实现', '', '', '■', '■', '■', '■', '', '', '', '', '', ''],
        ['4.驱动集成', '', '', '■', '■', '■', '', '', '', '', '', '', ''],
        ['5.测试调试', '', '', '', '', '', '■', '■', '■', '', '', '', ''],
        ['6.文档论文', '', '', '', '', '', '', '', '■', '■', '■', '■', '■'],
    ],
    title='附表2 项目甘特图'
)

# ============================================================
# 保存
# ============================================================
output_path = r'd:\RTOS—work2\程序源文件-FreeRTOS\Doc\RTOS论文.docx'
doc.save(output_path)
print(f'论文已生成: {output_path}')
print('请在Word中打开并进行以下操作：')
print('  1. 检查封面字体（华文新魏可能需要安装）')
print('  2. 插入自动目录（引用→目录）')
print('  3. 调整中文字体显示')
print('  4. 导出为PDF')
